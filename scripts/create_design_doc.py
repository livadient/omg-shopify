"""Generate the Google Ads API design document as .docx — v2 (production-ready emphasis)"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    return h


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


# Title
title = doc.add_heading("Google Ads API \u2014 Tool Design Document", 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

# Status banner
add_bold_para(
    "STATUS: PRODUCTION \u2014 This tool is fully developed, deployed, and running "
    "in production on an Azure VM. It has been operational since April 2026 and executes "
    "daily. We are requesting Basic access to switch from the test token to a live token "
    "so the tool can query real (non-test) keyword data."
)

# 1
add_heading("1. Tool Name")
doc.add_paragraph("OMG Shop Marketing Advisor (internal tool, codename: Atlas)")

# 2
add_heading("2. Company / Developer Name")
doc.add_paragraph(
    "OMG \u2014 a Cyprus-based e-commerce company operating online stores at "
    "omg.com.cy (Cyprus market) and omg.gr (Greece market). "
    "We sell custom graphic t-shirts and ship to Cyprus, Greece, and across Europe."
)

# 3
add_heading("3. Tool Overview")
doc.add_paragraph(
    "Atlas is a fully built and deployed internal marketing tool that runs as part of our "
    "e-commerce backend service. It is NOT a tool under development \u2014 it is live and operational."
)
doc.add_paragraph(
    "The tool runs on a daily automated schedule (Monday\u2013Friday at 07:00 Cyprus time) and "
    "generates marketing intelligence reports that are emailed to the store owners. "
    "It currently uses Google Search Console data (already integrated and working) and "
    "needs Google Ads Keyword Planner data to provide real CPC and search volume estimates "
    "instead of AI-generated guesses."
)

add_heading("What the tool does today (already in production):", level=2)
doc.add_paragraph("Fetches product catalog from our Shopify store via Admin API", style="List Bullet")
doc.add_paragraph("Fetches real search performance data from Google Search Console API (integrated, working)", style="List Bullet")
doc.add_paragraph("Fetches trending topics from Google Trends (integrated, working)", style="List Bullet")
doc.add_paragraph("Generates a daily marketing report with SEO and advertising recommendations", style="List Bullet")
doc.add_paragraph("Emails the report to the store owners for manual review and action", style="List Bullet")

add_heading("What we need Basic access for:", level=2)
doc.add_paragraph(
    "The only missing piece is real keyword data from the Keyword Planner. "
    "Currently the tool estimates keyword volumes and CPC using AI, which is inaccurate. "
    "With Basic access, we will call KeywordPlanIdeaService.GenerateKeywordIdeas to get "
    "real search volumes and CPC ranges, making the reports actionable."
)

# 4
add_heading("4. Google Ads API Usage \u2014 Specific and Limited")

doc.add_paragraph("We will use exactly ONE API service:")
add_bold_para("KeywordPlanIdeaService.GenerateKeywordIdeas")

doc.add_paragraph(
    "This service returns keyword suggestions with real monthly search volume, "
    "CPC estimates (low/high range), and competition level for a given geo target and language."
)

add_heading("We will NOT use:", level=2)
doc.add_paragraph("Campaign management (create, update, delete campaigns)", style="List Bullet")
doc.add_paragraph("Ad group management", style="List Bullet")
doc.add_paragraph("Bidding or budget management", style="List Bullet")
doc.add_paragraph("Reporting or analytics APIs", style="List Bullet")
doc.add_paragraph("Customer management", style="List Bullet")
doc.add_paragraph("Any write operations \u2014 our usage is 100% read-only", style="List Bullet")

# 5
add_heading("5. Detailed Technical Implementation")

doc.add_paragraph(
    "The implementation is complete. The source code is deployed and running. "
    "Below is the exact API call our tool makes:"
)

add_heading("API Call Details", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Service", "KeywordPlanIdeaService"],
        ["Method", "GenerateKeywordIdeas"],
        ["Customer ID", "9820211305"],
        ["Frequency", "1 call per day, Monday\u2013Friday"],
        ["Maximum calls per week", "5"],
        ["Results requested per call", "Up to 30 keyword ideas"],
        ["Authentication", "OAuth2 (Desktop app flow)"],
        ["Client library", "google-ads Python SDK (official)"],
    ],
)

add_heading("Geo Targets (rotated daily)", level=2)
add_table(
    ["Day", "Market", "Geo Target Constant", "Language Constant"],
    [
        ["Monday, Wednesday", "Cyprus", "geoTargetConstants/2196", "languageConstants/1022 (Greek)"],
        ["Tuesday, Thursday", "Greece", "geoTargetConstants/2300", "languageConstants/1022 (Greek)"],
        ["Friday", "Europe", "geoTargetConstants/2067", "languageConstants/1000 (English)"],
    ],
)

add_heading("Seed Keywords (example for Cyprus)", level=2)
doc.add_paragraph("t-shirt cyprus", style="List Bullet")
doc.add_paragraph("\u03bc\u03c0\u03bb\u03bf\u03c5\u03b6\u03ac\u03ba\u03b9\u03b1 \u03ba\u03cd\u03c0\u03c1\u03bf\u03c2", style="List Bullet")
doc.add_paragraph("graphic tees cyprus", style="List Bullet")
doc.add_paragraph("custom t-shirt", style="List Bullet")
doc.add_paragraph("funny t-shirt", style="List Bullet")
doc.add_paragraph("slogan tee", style="List Bullet")

add_heading("Example API Request (production code)", level=2)
add_code(
    "from google.ads.googleads.client import GoogleAdsClient\n\n"
    "client = GoogleAdsClient.load_from_dict(config)\n"
    "service = client.get_service('KeywordPlanIdeaService')\n"
    "request = client.get_type('GenerateKeywordIdeasRequest')\n"
    "request.customer_id = '9820211305'\n"
    "request.language = 'languageConstants/1022'\n"
    "request.geo_target_constants.append('geoTargetConstants/2196')\n"
    "request.keyword_seed.keywords.extend([\n"
    "    't-shirt cyprus', 'graphic tees', 'custom t-shirt'\n"
    "])\n"
    "response = service.generate_keyword_ideas(request=request)\n"
    "# Results included in daily email report"
)

# 6
add_heading("6. Production Deployment Details")
doc.add_paragraph(
    "The tool is deployed and running in production. It is not a prototype or proof-of-concept."
)
add_table(
    ["Aspect", "Detail"],
    [
        ["Server", "Azure VM (Ubuntu Linux), Docker container"],
        ["Runtime", "Python 3.13, FastAPI web framework"],
        ["Uptime", "24/7, auto-restart on failure"],
        ["Schedule", "APScheduler cron job, Mon\u2013Fri 07:00 Europe/Nicosia"],
        ["Deployment", "Automated via GitHub Actions CI/CD pipeline"],
        ["Live since", "April 2026"],
        ["Other integrations already working", "Google Search Console API, Google Trends, Shopify Admin API"],
    ],
)

# 7
add_heading("7. Output: Daily Email Report")
doc.add_paragraph(
    "The tool sends a daily email report to the store owners. "
    "The Keyword Planner data appears in the 'Google Ads Suggestions' section. "
    "Here is what the output looks like:"
)
add_table(
    ["Keyword", "Est. CPC", "Volume", "Competition", "Note"],
    [
        ["custom t-shirts cyprus", "EUR 0.30", "320/mo", "LOW", "High purchase intent"],
        ["graphic tees limassol", "EUR 0.15", "50/mo", "LOW", "Local targeting"],
        ["\u03bc\u03c0\u03bb\u03bf\u03c5\u03b6\u03ac\u03ba\u03b9\u03b1 online", "EUR 0.20", "180/mo", "MEDIUM", "Greek language market"],
    ],
)
doc.add_paragraph(
    "A human (store owner) reviews the report and manually decides which keywords "
    "to target. The tool does NOT create or modify any campaigns automatically."
)

# 8
add_heading("8. OAuth2 Implementation")
add_table(
    ["Aspect", "Detail"],
    [
        ["Authentication type", "OAuth2 (Desktop application flow)"],
        ["Client type", "Desktop app (created in Google Cloud Console)"],
        ["Scope", "https://www.googleapis.com/auth/adwords"],
        ["Token storage", "Refresh token in server environment variables (.env)"],
        ["Token refresh", "Automatic via google-ads Python client library"],
        ["Google Cloud Project", "omg-shop-492712"],
    ],
)

# 9
add_heading("9. Data Handling and Privacy")
doc.add_paragraph("Keyword data is used only in internal email reports to store owners", style="List Bullet")
doc.add_paragraph("No keyword data is stored permanently (only included in the email)", style="List Bullet")
doc.add_paragraph("No data is shared with any third party", style="List Bullet")
doc.add_paragraph("No end-user or customer data is sent to or from the Google Ads API", style="List Bullet")
doc.add_paragraph("The tool is internal-use only \u2014 not a SaaS product, not resold, no external users", style="List Bullet")

# 10
add_heading("10. Contact Information")
add_table(
    ["Field", "Value"],
    [
        ["Developer", "Vangelis Livadiotis"],
        ["Email", "livadient@gmail.com"],
        ["Company website", "https://omg.com.cy"],
        ["Google Ads Customer ID", "982-021-1305"],
        ["Google Cloud Project ID", "omg-shop-492712"],
    ],
)

out = "doc/google-ads-api-design-doc.docx"
doc.save(out)
print(f"Created: {out}")
